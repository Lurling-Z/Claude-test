#!/usr/bin/env python3
"""将雅思自学路径系列 5 篇仿写内容导出为 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

with open('/projects/sandbox/Claude-test/content/2026-summer-ielts-self-study-xhs.md', 'r', encoding='utf-8') as f:
    content = f.read()

# 标题页
title = doc.add_heading('渊学通杭州', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('雅思自学路径系列 · 小红书仿写 5 篇', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('账号：渊学通杭州（小红书认证专业号）\n基于 5 篇雅思赛道高赞爆款做原创仿写')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
compliance = doc.add_paragraph()
compliance.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = compliance.add_run('套路提炼：')
run.bold = True
run.font.size = Pt(9)
run = compliance.add_run('省钱感 · 免费资料 · 路线图 · 低门槛 · 评论领取')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
compliance = doc.add_paragraph()
compliance.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = compliance.add_run('合规基线：')
run.bold = True
run.font.size = Pt(9)
run = compliance.add_run('遵循小红书《社区公约 2.0》（2026.1）。规避"保过 / 包出分 / 7 天逆袭 / 必上岸 / 8 分易如反掌 / 一个月冲 8 / 慈善家"等强营销话术。')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# 按选题拆分
topics = content.split('## 选题')[1:]

for i, topic in enumerate(topics):
    if '发布顺序建议' in topic or '合规自查清单' in topic:
        # 处理后续章节
        # 发布顺序建议
        if '## 发布顺序建议' in topic:
            doc.add_heading('发布顺序建议', level=1)
            order_data = [
                ('第 1 周 · 周一', '选题 1（资料）', '吃公域流量，账号曝光起量'),
                ('第 1 周 · 周四', '选题 5（零基础诊断）', '把流量接到诊断 / 课程咨询'),
                ('第 2 周 · 周一', '选题 2（30 天路线图）', '高收藏型，账号沉淀'),
                ('第 2 周 · 周四', '选题 3（口语串题）', '垂直深度，建立专业信任'),
                ('第 2 周 · 周日', '选题 4（写作自查）', '同上，覆盖写作人群'),
            ]
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            header_cells = table.rows[0].cells
            header_cells[0].text = '顺序'
            header_cells[1].text = '选题'
            header_cells[2].text = '目的'
            for time_slot, topic_name, purpose in order_data:
                row_cells = table.add_row().cells
                row_cells[0].text = time_slot
                row_cells[1].text = topic_name
                row_cells[2].text = purpose

            doc.add_paragraph()

        # 合规清单
        doc.add_heading('合规自查清单（发布前过一遍）', level=1)
        checklist_items = [
            '全文未出现：保过 / 包出分 / 必上岸 / 7 天逆袭 / 8 分易如反掌 / 一个月 X 分 / 名师押题 / 命题组',
            '评论区领取话术不带"包过 / 一定能 / 免费送几万资料"等措辞',
            '涉及对比时未点名贬低同行',
            '涉及价格时未做夸张低价引流',
            'AI 辅助创作（如有）已在文末或图中标注',
            '未涉及未成年人具体身份信息（姓名 / 学校 / 成绩单截图露脸）',
            '评论区领取需通过私信发资料，避免在评论里直接放外链',
        ]
        for item in checklist_items:
            p = doc.add_paragraph(f'☐ {item}')
            p.paragraph_format.space_after = Pt(4)
        break

    # 解析标题
    lines = topic.strip().split('\n')
    first_line = lines[0].strip()
    topic_title = f'选题 {first_line.split("｜")[0].strip()}' if '｜' in first_line else f'选题 {i+1}'
    topic_subtitle = first_line.split('｜', 1)[1].strip() if '｜' in first_line else first_line

    doc.add_heading(f'{topic_title}｜{topic_subtitle}', level=1)

    body_started = False

    for line in lines[1:]:
        line = line.strip()

        if not line:
            if body_started:
                doc.add_paragraph()
            continue

        if line == '---':
            continue

        # 元信息
        if line.startswith('- **定位**'):
            p = doc.add_paragraph()
            run = p.add_run('定位：')
            run.bold = True
            run.font.size = Pt(10)
            value = line.replace('- **定位**：', '').strip()
            run = p.add_run(value)
            run.font.size = Pt(10)
            continue

        if line.startswith('- **封面文案**'):
            p = doc.add_paragraph()
            run = p.add_run('封面文案：')
            run.bold = True
            run.font.size = Pt(10)
            value = line.replace('- **封面文案**：', '').strip()
            run = p.add_run(value)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(200, 50, 50)
            continue

        if line.startswith('- **标题**'):
            p = doc.add_paragraph()
            run = p.add_run('标题：')
            run.bold = True
            run.font.size = Pt(10)
            value = line.replace('- **标题**：', '').strip()
            run = p.add_run(value)
            run.font.size = Pt(10)
            continue

        if line == '### 正文':
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run('—— 正文 ——')
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)
            body_started = True
            continue

        if line.startswith('**标签**'):
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run('标签：')
            run.bold = True
            run.font.size = Pt(9)
            tags = line.replace('**标签**：', '').strip()
            run = p.add_run(tags)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(50, 100, 200)
            continue

        # 正文内容
        if body_started:
            clean_line = line.replace('**', '')
            emoji_starts = ['1️⃣', '2️⃣', '3️⃣', '4️⃣', '5️⃣', '⚠️', '✅', '❎', '📌', '📅',
                            '🌅', '🍳', '📚', '🍱', '✍️', '🏃', '🧠', '🌙', '📵', '📊',
                            '💰', '👥', '👩', '🏫', '🔍']
            if any(line.startswith(e) for e in emoji_starts):
                p = doc.add_paragraph()
                run = p.add_run(clean_line)
                run.bold = True
                run.font.size = Pt(11)
            elif line.startswith('- ') or line.startswith('· '):
                p = doc.add_paragraph(clean_line.lstrip('-· '), style='List Bullet')
                p.paragraph_format.space_after = Pt(2)
            else:
                p = doc.add_paragraph(clean_line)
                p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

output_path = '/projects/sandbox/Claude-test/content/渊学通杭州_雅思自学路径系列_小红书5篇.docx'
doc.save(output_path)
print(f'✅ Word 文件已生成：{output_path}')
