#!/usr/bin/env python3
"""将 9 宫格图文脚本导出为 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

with open('/projects/sandbox/Claude-test/content/2026-summer-ielts-camp-9grid.md', 'r', encoding='utf-8') as f:
    content = f.read()

# 标题页
title = doc.add_heading('渊学通杭州', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('2026 雅思暑期封闭营 · 9 宫格图文脚本', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('共 5 组 × 9 张 = 45 张图\n含视觉规范 + 文案 + 设计方向')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# 按行解析
lines = content.split('\n')
i = 0
while i < len(lines):
    line = lines[i].strip()

    # 跳过空行
    if not line:
        i += 1
        continue

    # 一级标题 ##
    if line.startswith('## ') and not line.startswith('### '):
        heading_text = line.replace('## ', '').strip()
        if heading_text.startswith('视觉总规范'):
            doc.add_heading('视觉总规范（5 篇通用）', level=1)
        elif heading_text.startswith('5 篇通用'):
            doc.add_heading('5 篇通用：账号视觉一致性建议', level=1)
        elif heading_text.startswith('发布前自查'):
            doc.add_heading('发布前自查清单', level=1)
        else:
            doc.add_heading(heading_text, level=1)
        i += 1
        continue

    # 三级标题 ### (P1-P9)
    if line.startswith('### '):
        heading_text = line.replace('### ', '').strip()
        doc.add_heading(heading_text, level=2)
        i += 1
        continue

    # 表格行
    if line.startswith('|'):
        # 简单处理表格为文本
        if '---' in line:
            i += 1
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        p = doc.add_paragraph()
        for j, cell in enumerate(cells):
            run = p.add_run(cell)
            if j == 0:
                run.bold = True
            if j < len(cells) - 1:
                p.add_run('　｜　')
        p.paragraph_format.space_after = Pt(2)
        i += 1
        continue

    # 加粗行
    if line.startswith('**') and line.endswith('**'):
        clean = line.strip('*').strip()
        p = doc.add_paragraph()
        run = p.add_run(clean)
        run.bold = True
        i += 1
        continue

    # 主文案/副文案/视觉标记
    if line.startswith('**主文案**'):
        p = doc.add_paragraph()
        run = p.add_run('主文案：')
        run.bold = True
        run.font.color.rgb = RGBColor(200, 50, 50)
        value = line.replace('**主文案**：', '').replace('**主文案**:', '').strip()
        run = p.add_run(value)
        run.font.size = Pt(12)
        run.bold = True
        i += 1
        continue

    if line.startswith('**副文案**'):
        p = doc.add_paragraph()
        run = p.add_run('副文案：')
        run.bold = True
        run.font.color.rgb = RGBColor(50, 100, 150)
        value = line.replace('**副文案**：', '').replace('**副文案**:', '').strip()
        run = p.add_run(value)
        run.font.size = Pt(10)
        i += 1
        continue

    if line.startswith('**视觉**'):
        p = doc.add_paragraph()
        run = p.add_run('视觉建议：')
        run.bold = True
        run.font.color.rgb = RGBColor(50, 150, 50)
        value = line.replace('**视觉**：', '').replace('**视觉**:', '').strip()
        run = p.add_run(value)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)
        i += 1
        continue

    # 分隔线
    if line == '---':
        doc.add_paragraph('─' * 40)
        doc.add_page_break()
        i += 1
        continue

    # 列表
    if line.startswith('- [ ]') or line.startswith('- [x]'):
        clean = line.replace('- [ ]', '☐').replace('- [x]', '☑')
        p = doc.add_paragraph(clean)
        p.paragraph_format.space_after = Pt(2)
        i += 1
        continue

    if line.startswith('- ') or line.startswith('· '):
        clean = line.lstrip('- ·').strip()
        p = doc.add_paragraph(f'• {clean}')
        p.paragraph_format.space_after = Pt(2)
        i += 1
        continue

    # 引用块
    if line.startswith('>'):
        clean = line.lstrip('> ').strip()
        p = doc.add_paragraph()
        run = p.add_run(clean)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.italic = True
        i += 1
        continue

    # 普通段落
    clean = line.replace('**', '')
    if clean:
        p = doc.add_paragraph(clean)
        p.paragraph_format.space_after = Pt(4)

    i += 1

output_path = '/projects/sandbox/Claude-test/content/渊学通杭州_2026雅思暑期封闭营_9宫格脚本.docx'
doc.save(output_path)
print(f'✅ Word 文件已生成：{output_path}')
