#!/usr/bin/env python3
"""将小红书选题内容导出为 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# 读取 Markdown 文件
with open('/projects/sandbox/Claude-test/content/2026-summer-ielts-camp-xhs.md', 'r', encoding='utf-8') as f:
    content = f.read()

# 标题页
title = doc.add_heading('渊学通杭州', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('2026 雅思暑期封闭营 · 小红书选题 5 篇', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('账号：渊学通杭州（小红书认证专业号）\n选题主线：雅思封闭营 / 国际生暑期规划')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
compliance = doc.add_paragraph()
compliance.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = compliance.add_run('合规基线：')
run.bold = True
run.font.size = Pt(9)
run = compliance.add_run('遵循小红书《社区公约 2.0》（2026.1）。规避"保过 / 包出分 / 7 天逆袭 / 必上岸 / 保证提分 / 名师押题 / 命题组老师 / 焦虑式话术"等表达。整体调性：专业判断 + 真实经验，口语化、不堆卖点。')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# 按选题拆分
topics = content.split('## 选题')[1:]  # 跳过第一段（文件头）

for i, topic in enumerate(topics):
    if '合规自查清单' in topic:
        # 最后一段是合规清单
        doc.add_heading('合规自查清单（发布前过一遍）', level=1)
        checklist_items = [
            '全文未出现：保过 / 包出分 / 必上岸 / 保证提分 / 7 天逆袭 / 名师押题 / 命题组',
            '全文未出现"不学就落后 / 别人都在学"等焦虑话术',
            '提分、出分均使用"适合做集中训练 / 稳定区间 / 多数学员的反馈"等表述',
            '涉及对比时，未点名贬低同行',
            '涉及价格时，未做夸张低价引流，仅讲价格构成',
            'AI 辅助创作（如有）已在文末或图中标注',
            '未涉及未成年人具体身份信息（姓名 / 学校 / 成绩单截图露脸）',
        ]
        for item in checklist_items:
            p = doc.add_paragraph(f'☐ {item}')
            p.paragraph_format.space_after = Pt(4)
        break

    # 解析标题
    lines = topic.strip().split('\n')
    first_line = lines[0].strip()
    # 选题标题
    topic_title = f'选题 {first_line.split("｜")[0].strip()}' if '｜' in first_line else f'选题 {i+1}'
    topic_subtitle = first_line.split('｜')[1].strip() if '｜' in first_line else first_line

    doc.add_heading(f'{topic_title}｜{topic_subtitle}', level=1)

    # 解析定位、封面文案、标题
    meta_section = True
    body_started = False

    for line in lines[1:]:
        line = line.strip()

        if not line:
            if body_started:
                doc.add_paragraph()
            continue

        if line == '---':
            continue

        # 元信息行
        if line.startswith('- **定位**'):
            p = doc.add_paragraph()
            run = p.add_run('定位：')
            run.bold = True
            run.font.size = Pt(10)
            value = line.replace('- **定位**：', '').strip()
            run = p.add_run(value)
            run.font.size = Pt(10)
            continue

        if line.startswith('- **封面文案**'):
            p = doc.add_paragraph()
            run = p.add_run('封面文案：')
            run.bold = True
            run.font.size = Pt(10)
            value = line.replace('- **封面文案**：', '').strip()
            run = p.add_run(value)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(200, 50, 50)
            continue

        if line.startswith('- **标题**'):
            p = doc.add_paragraph()
            run = p.add_run('标题：')
            run.bold = True
            run.font.size = Pt(10)
            value = line.replace('- **标题**：', '').replace('- **标题候选**', '').strip()
            run = p.add_run(value)
            run.font.size = Pt(10)
            continue

        if line == '### 正文':
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run('—— 正文 ——')
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)
            body_started = True
            continue

        if line.startswith('**标签**'):
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run('标签：')
            run.bold = True
            run.font.size = Pt(9)
            tags = line.replace('**标签**：', '').strip()
            run = p.add_run(tags)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(50, 100, 200)
            continue

        # 正文内容
        if body_started:
            # 处理加粗
            clean_line = line.replace('**', '')
            # 处理 emoji 标题行（如 1️⃣）
            if any(line.startswith(e) for e in ['1️⃣', '2️⃣', '3️⃣', '4️⃣', '5️⃣', '⚠️', '✅', '❎', '📌', '🌅', '🍳', '📚', '🍱', '✍️', '🏃', '🧠', '🌙', '📵', '📊', '💰', '👥', '👩', '🏫', '🔍']):
                p = doc.add_paragraph()
                run = p.add_run(clean_line)
                run.bold = True
                run.font.size = Pt(11)
            elif line.startswith('- ') or line.startswith('· '):
                p = doc.add_paragraph(clean_line, style='List Bullet')
                p.paragraph_format.space_after = Pt(2)
            else:
                p = doc.add_paragraph(clean_line)
                p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

# 保存
output_path = '/projects/sandbox/Claude-test/content/渊学通杭州_2026雅思暑期封闭营_小红书5篇.docx'
doc.save(output_path)
print(f'✅ Word 文件已生成：{output_path}')
